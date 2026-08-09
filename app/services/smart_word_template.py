from __future__ import annotations

from collections import defaultdict
from copy import deepcopy
from datetime import date, time, timedelta
from hashlib import sha256
from io import BytesIO
from pathlib import Path
from typing import Iterable
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.text import WD_BREAK

from app.services.word_export import WordExportError
from app.utils import as_local


TEMPLATE_SHA256 = "e32eb2abea792f23a9ba86c7c69949c36b67f2544c1a309d590e7d57512b2a69"
TEMPLATE_PATH = Path(__file__).resolve().parents[1] / "assets" / "amuda_schedule_template.docx"
SLOTS_PER_SIDE = 19
DATES_PER_PAGE = SLOTS_PER_SIDE * 2

LEVANT_MONTHS = {
    1: "كانون الثاني",
    2: "شباط",
    3: "آذار",
    4: "نيسان",
    5: "أيار",
    6: "حزيران",
    7: "تموز",
    8: "آب",
    9: "أيلول",
    10: "تشرين الأول",
    11: "تشرين الثاني",
    12: "كانون الأول",
}
ARABIC_DAYS = {
    0: "الاثنين",
    1: "الثلاثاء",
    2: "الأربعاء",
    3: "الخميس",
    4: "الجمعة",
    5: "السبت",
    6: "الاحد",
}


def _template_bytes() -> bytes:
    try:
        data = TEMPLATE_PATH.read_bytes()
    except OSError as exc:
        raise WordExportError("قالب Word الرسمي غير موجود على الخادم.") from exc
    if sha256(data).hexdigest() != TEMPLATE_SHA256:
        raise WordExportError("قالب Word الرسمي تغيّر أو تلف، لذلك تم إيقاف التصدير حفاظاً على التنسيق.")
    return data


def _group_shifts(shifts: Iterable, timezone: ZoneInfo) -> dict[date, dict[str, list[str]]]:
    grouped: dict[date, dict[str, list[str]]] = defaultdict(lambda: {"day": [], "evening": []})
    for shift in sorted(shifts, key=lambda item: (item.start_at, item.pharmacy.name)):
        if not getattr(shift, "active", True):
            continue
        local_start = as_local(shift.start_at, timezone)
        period = "day" if local_start.time().replace(tzinfo=None) < time(18, 0) else "evening"
        name = shift.pharmacy.name.strip()
        if name and name not in grouped[local_start.date()][period]:
            grouped[local_start.date()][period].append(name)
    return dict(grouped)


def _expected_dates(
    grouped: dict[date, dict[str, list[str]]],
    *,
    period_start: date | None,
    period_end: date | None,
) -> list[date]:
    if period_start is None and period_end is None:
        return sorted(grouped)
    if period_start is None or period_end is None or period_end < period_start:
        raise WordExportError("فترة الجدول غير صالحة لتصدير Word.")
    dates: list[date] = []
    current = period_start
    while current <= period_end:
        dates.append(current)
        current += timedelta(days=1)
    return dates


def _validate_complete(
    grouped: dict[date, dict[str, list[str]]],
    dates: list[date],
) -> None:
    if not dates:
        raise WordExportError("لا توجد مناوبات قابلة للتصدير إلى Word.")
    for duty_date in dates:
        values = grouped.get(duty_date, {"day": [], "evening": []})
        if len(values["day"]) != 1 or len(values["evening"]) != 1:
            raise WordExportError(
                "لا يمكن إنشاء Word لأن بيانات يوم "
                f"{duty_date.strftime('%d/%m/%Y')} لا تحتوي صيدلية نهارية وصيدلية مسائية واحدة بالضبط."
            )


def _replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        first = paragraph.runs[0]
        first.text = text
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


def _replace_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    _replace_paragraph_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def _date_label(value: date) -> str:
    return f"{ARABIC_DAYS[value.weekday()]}: {value.day}/{value.month}/{value.year}"


def _title_for(value: date) -> str:
    return (
        "جدول المنوبات لمدينة عامودة خلال "
        f"شهر//{value.month}// {LEVANT_MONTHS[value.month]} لعام {value.year}"
    )


def _validate_template_structure(document) -> None:
    if len(document.tables) != 1:
        raise WordExportError("قالب Word الرسمي لا يحتوي الجدول المتوقع.")
    table = document.tables[0]
    if len(table.rows) != 20 or len(table.columns) != 6:
        raise WordExportError("بنية جدول Word الرسمي تغيّرت، لذلك تم إيقاف التصدير حفاظاً على القالب.")
    titles = [paragraph for paragraph in document.paragraphs if paragraph.text.startswith("جدول المنوبات")]
    footers = [paragraph for paragraph in document.paragraphs if paragraph.text.startswith("الدوام المسائي")]
    if len(titles) != 1 or len(footers) != 1:
        raise WordExportError("عناوين قالب Word الرسمي غير مطابقة للقالب المعتمد.")


def _append_template_page(document, template) -> None:
    body = document.element.body
    section_properties = body.sectPr
    page_break = document.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)
    # Template body order is: blank paragraph, title, table, footer, section properties.
    for child in (template.element.body[1], template.element.body[2], template.element.body[3]):
        section_properties.addprevious(deepcopy(child))


def _fill_page(table, page_dates: list[date], grouped: dict[date, dict[str, list[str]]]) -> None:
    for slot in range(DATES_PER_PAGE):
        row_index = 1 + (slot % SLOTS_PER_SIDE)
        base = 0 if slot < SLOTS_PER_SIDE else 3
        if slot >= len(page_dates):
            _replace_cell_text(table.rows[row_index].cells[base], "")
            _replace_cell_text(table.rows[row_index].cells[base + 1], "")
            _replace_cell_text(table.rows[row_index].cells[base + 2], "")
            continue

        duty_date = page_dates[slot]
        values = grouped[duty_date]
        _replace_cell_text(table.rows[row_index].cells[base], _date_label(duty_date))
        _replace_cell_text(table.rows[row_index].cells[base + 1], values["day"][0])
        _replace_cell_text(table.rows[row_index].cells[base + 2], values["evening"][0])


def build_smart_template_schedule(
    shifts: Iterable,
    timezone: ZoneInfo,
    *,
    period_start: date | None = None,
    period_end: date | None = None,
) -> bytes:
    """Fill the user's official Word template from the smart-schedule rows.

    The template's table design, headers, colors, fonts, row sizing and footer are
    preserved. Only schedule content (date labels/title and the two pharmacy-name
    cells for each day) is filled from the current smart draft/published batch.
    """
    grouped = _group_shifts(shifts, timezone)
    dates = _expected_dates(grouped, period_start=period_start, period_end=period_end)
    _validate_complete(grouped, dates)

    template_data = _template_bytes()
    document = Document(BytesIO(template_data))
    _validate_template_structure(document)
    template = Document(BytesIO(template_data))

    pages = [dates[index : index + DATES_PER_PAGE] for index in range(0, len(dates), DATES_PER_PAGE)]
    for _ in range(1, len(pages)):
        _append_template_page(document, template)

    tables = document.tables
    titles = [paragraph for paragraph in document.paragraphs if paragraph.text.startswith("جدول المنوبات")]
    if len(tables) != len(pages) or len(titles) != len(pages):
        raise WordExportError("تعذر تكرار قالب Word لكل صفحات الجدول.")

    for page_index, page_dates in enumerate(pages):
        _replace_paragraph_text(titles[page_index], _title_for(page_dates[0]))
        _fill_page(tables[page_index], page_dates, grouped)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
