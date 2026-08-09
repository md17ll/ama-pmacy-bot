from __future__ import annotations

from collections import defaultdict
from copy import deepcopy
from datetime import date, time
from io import BytesIO
from pathlib import Path
from typing import Iterable
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.text import WD_BREAK
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.models import Shift
from app.services.shift_schedule_tools import ShiftTimes
from app.utils import as_local


MAX_DATES_PER_PAGE = 38
ROWS_PER_SIDE = 19
TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "templates" / "official_schedule_template.docx"
LEVANT_MONTHS = {
    1: "كانون الثاني",
    2: "شباط",
    3: "آذار",
    4: "نيسان",
    5: "أيار",
    6: "حزيران",
    7: "تموز",
    8: "آب",
    9: "أيلول",
    10: "تشرين الأول",
    11: "تشرين الثاني",
    12: "كانون الأول",
}
ARABIC_DAYS = {
    0: "الاثنين",
    1: "الثلاثاء",
    2: "الأربعاء",
    3: "الخميس",
    4: "الجمعة",
    5: "السبت",
    6: "الاحد",
}


class WordExportError(ValueError):
    pass


def _load_template():
    if not TEMPLATE_PATH.is_file():
        raise WordExportError("قالب Word الرسمي غير موجود داخل النظام.")
    document = Document(str(TEMPLATE_PATH))
    if len(document.tables) != 1:
        raise WordExportError("قالب Word الرسمي غير صالح: يجب أن يحتوي جدولاً واحداً.")
    table = document.tables[0]
    if len(table.rows) != ROWS_PER_SIDE + 1 or len(table.columns) != 6:
        raise WordExportError("قالب Word الرسمي غير صالح: بنية الجدول تغيرت.")
    title = next(
        (paragraph for paragraph in document.paragraphs if "جدول المنوبات لمدينة عامودة" in paragraph.text),
        None,
    )
    footer = next(
        (paragraph for paragraph in document.paragraphs if "الدوام المسائي" in paragraph.text),
        None,
    )
    if title is None or footer is None:
        raise WordExportError("قالب Word الرسمي غير صالح: العنوان أو التذييل غير موجود.")
    return document, title, table, footer


def _replace_paragraph_text(paragraph, text: str, *, style_run=None) -> None:
    """Replace visible text while preserving the exact formatting of the template."""
    if paragraph.runs:
        run = paragraph.runs[0]
        run.text = text
        if style_run is not None and run._r.rPr is None and style_run._r.rPr is not None:
            run._r.insert(0, deepcopy(style_run._r.rPr))
        for extra in paragraph.runs[1:]:
            extra.text = ""
        return

    run = paragraph.add_run(text)
    if style_run is not None and style_run._r.rPr is not None:
        run._r.insert(0, deepcopy(style_run._r.rPr))


def _first_styled_run(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.text.strip() or run._r.rPr is not None:
                return run
    return None


def _replace_cell_text(cell, text: str, *, style_cell=None) -> None:
    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = cell.paragraphs[0]
    style_run = _first_styled_run(style_cell) if style_cell is not None else None
    _replace_paragraph_text(paragraph, text, style_run=style_run)
    for extra in cell.paragraphs[1:]:
        for run in extra.runs:
            run.text = ""


def _date_label(value: date) -> str:
    return f"{ARABIC_DAYS[value.weekday()]}: {value.day}/{value.month}/{value.year}"


def _title_for(first_date: date) -> str:
    return (
        "جدول المنوبات لمدينة عامودة خلال "
        f"شهر//{first_date.month}// {LEVANT_MONTHS[first_date.month]} لعام {first_date.year}"
    )


def _group_shifts(shifts: Iterable[Shift], timezone: ZoneInfo) -> dict[date, dict[str, list[str]]]:
    grouped: dict[date, dict[str, list[str]]] = defaultdict(lambda: {"day": [], "evening": []})
    for shift in sorted(shifts, key=lambda item: (item.start_at, item.pharmacy.name)):
        if not getattr(shift, "active", True):
            continue
        local_start = as_local(shift.start_at, timezone)
        period = "day" if local_start.time().replace(tzinfo=None) < time(18, 0) else "evening"
        name = shift.pharmacy.name.strip()
        if name and name not in grouped[local_start.date()][period]:
            grouped[local_start.date()][period].append(name)
    return dict(grouped)


def _joined_names(values: list[str]) -> str:
    return " / ".join(values)


def _fill_page(
    title: Paragraph,
    table: Table,
    page_dates: list[date],
    grouped: dict[date, dict[str, list[str]]],
) -> None:
    _replace_paragraph_text(title, _title_for(page_dates[0]))

    style_cells = {
        0: table.rows[1].cells[0],
        1: table.rows[1].cells[1],
        2: table.rows[1].cells[2],
        3: table.rows[1].cells[3],
        4: table.rows[1].cells[4],
        5: table.rows[1].cells[5],
    }

    for row_index in range(ROWS_PER_SIDE):
        row = table.rows[row_index + 1]
        for base, date_index in ((0, row_index), (3, ROWS_PER_SIDE + row_index)):
            if date_index >= len(page_dates):
                for offset in range(3):
                    _replace_cell_text(row.cells[base + offset], "")
                continue

            duty_date = page_dates[date_index]
            values = grouped[duty_date]
            _replace_cell_text(
                row.cells[base],
                _date_label(duty_date),
                style_cell=style_cells[base],
            )
            _replace_cell_text(
                row.cells[base + 1],
                _joined_names(values["day"]),
                style_cell=style_cells[base + 1],
            )
            _replace_cell_text(
                row.cells[base + 2],
                _joined_names(values["evening"]),
                style_cell=style_cells[base + 2],
            )


def _append_template_page(document, pristine, page_dates, grouped) -> None:
    body = document._element.body
    section_properties = body.sectPr

    page_break = document.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)

    title_element = deepcopy(pristine["title"])
    table_element = deepcopy(pristine["table"])
    footer_element = deepcopy(pristine["footer"])
    section_properties.addprevious(title_element)
    section_properties.addprevious(table_element)
    section_properties.addprevious(footer_element)

    title = Paragraph(title_element, document._body)
    table = Table(table_element, document._body)
    _fill_page(title, table, page_dates, grouped)


def build_official_word_schedule(
    shifts: Iterable[Shift],
    timezone: ZoneInfo,
    times: ShiftTimes,
) -> bytes:
    """Fill the user's uploaded Word layout with the smart-scheduler result.

    The uploaded DOCX is the authoritative visual template. The scheduler remains
    the authoritative source for pharmacy names and day/evening placement. Header,
    footer, dimensions, borders, fonts and other visual formatting come directly
    from the template and are not regenerated in code.
    """
    del times  # Times stay exactly as written in the approved Word template.

    grouped = _group_shifts(shifts, timezone)
    dates = sorted(grouped)
    if not dates:
        raise WordExportError("لا توجد مناوبات قابلة للتصدير إلى Word.")

    document, title, table, _footer = _load_template()
    _pristine_document, pristine_title, pristine_table, pristine_footer = _load_template()
    pristine = {
        "title": pristine_title._p,
        "table": pristine_table._tbl,
        "footer": pristine_footer._p,
    }

    _fill_page(title, table, dates[:MAX_DATES_PER_PAGE], grouped)

    for start in range(MAX_DATES_PER_PAGE, len(dates), MAX_DATES_PER_PAGE):
        _append_template_page(
            document,
            pristine,
            dates[start : start + MAX_DATES_PER_PAGE],
            grouped,
        )

    output = BytesIO()
    document.save(output)
    return output.getvalue()
