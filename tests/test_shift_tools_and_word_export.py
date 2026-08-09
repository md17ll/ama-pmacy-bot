from __future__ import annotations

import asyncio
from datetime import date, datetime, time, timezone
from io import BytesIO
from types import SimpleNamespace
from zoneinfo import ZoneInfo

from docx import Document
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.models import Base, Pharmacy, Shift
from app.services.shift_schedule_tools import (
    ShiftTimes,
    bulk_update_shift_times,
    get_shift_times,
    undo_bulk_time_update,
)
from app.services.word_export import build_official_word_schedule
from app.services.word_schedule import parse_amuda_word_schedule
from app.utils import as_local


TZ = ZoneInfo("Asia/Damascus")


def _utc(day: date, value: time) -> datetime:
    return datetime.combine(day, value, tzinfo=TZ).astimezone(timezone.utc)


class _AsyncSessionWrapper:
    def __init__(self, session):
        self._session = session

    async def __aenter__(self):
        return self

    async def __aexit__(self, exc_type, exc, tb):
        self._session.close()

    def add(self, value):
        self._session.add(value)

    def add_all(self, values):
        self._session.add_all(values)

    async def get(self, *args, **kwargs):
        return self._session.get(*args, **kwargs)

    async def scalars(self, *args, **kwargs):
        return self._session.scalars(*args, **kwargs)

    async def flush(self):
        self._session.flush()

    async def commit(self):
        self._session.commit()

    async def rollback(self):
        self._session.rollback()


def test_bulk_time_update_preserves_pharmacy_and_date_and_can_undo() -> None:
    async def run() -> None:
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        sessions = sessionmaker(engine, expire_on_commit=False)

        async with _AsyncSessionWrapper(sessions()) as session:
            pharmacy = Pharmacy(
                name="صيدلية النور",
                normalized_name="النور",
                address="شارع البلدية",
                status="active",
            )
            session.add(pharmacy)
            await session.flush()
            past = Shift(
                pharmacy_id=pharmacy.id,
                start_at=_utc(date(2026, 8, 4), time(13, 30)),
                end_at=_utc(date(2026, 8, 4), time(17, 0)),
                created_by=1,
                active=True,
            )
            day_shift = Shift(
                pharmacy_id=pharmacy.id,
                start_at=_utc(date(2026, 8, 10), time(13, 30)),
                end_at=_utc(date(2026, 8, 10), time(17, 0)),
                created_by=1,
                active=True,
            )
            evening_shift = Shift(
                pharmacy_id=pharmacy.id,
                start_at=_utc(date(2026, 8, 11), time(20, 30)),
                end_at=_utc(date(2026, 8, 11), time(23, 30)),
                created_by=1,
                active=True,
            )
            session.add_all([past, day_shift, evening_shift])
            await session.commit()
            original_day = (day_shift.start_at, day_shift.end_at)
            original_evening = (evening_shift.start_at, evening_shift.end_at)
            original_past = (past.start_at, past.end_at)
            pharmacy_id = pharmacy.id
            day_id = day_shift.id
            evening_id = evening_shift.id
            past_id = past.id

        new_times = ShiftTimes(
            day_start=time(14, 0),
            day_end=time(17, 30),
            evening_start=time(21, 0),
            evening_end=time(0, 0),
        )
        async with _AsyncSessionWrapper(sessions()) as session:
            count, audit_id = await bulk_update_shift_times(
                session,
                times=new_times,
                effective_at=_utc(date(2026, 8, 10), time.min),
                timezone=TZ,
                admin_id=99,
            )
            assert count == 2

        async with _AsyncSessionWrapper(sessions()) as session:
            updated_day = await session.get(Shift, day_id)
            updated_evening = await session.get(Shift, evening_id)
            unchanged_past = await session.get(Shift, past_id)
            assert updated_day is not None and updated_evening is not None and unchanged_past is not None
            assert updated_day.pharmacy_id == pharmacy_id
            assert updated_evening.pharmacy_id == pharmacy_id
            assert as_local(updated_day.start_at, TZ).date() == date(2026, 8, 10)
            assert as_local(updated_day.start_at, TZ).time() == time(14, 0)
            assert as_local(updated_day.end_at, TZ).time() == time(17, 30)
            assert as_local(updated_evening.start_at, TZ).date() == date(2026, 8, 11)
            assert as_local(updated_evening.start_at, TZ).time() == time(21, 0)
            assert as_local(updated_evening.end_at, TZ).date() == date(2026, 8, 12)
            assert as_local(updated_evening.end_at, TZ).time() == time(0, 0)
            assert as_local(unchanged_past.start_at, TZ) == as_local(original_past[0], TZ)
            assert as_local(unchanged_past.end_at, TZ) == as_local(original_past[1], TZ)
            stored = await get_shift_times(session)
            assert stored == new_times

        async with _AsyncSessionWrapper(sessions()) as session:
            restored = await undo_bulk_time_update(session, audit_id=audit_id, admin_id=99)
            assert restored == 2

        async with _AsyncSessionWrapper(sessions()) as session:
            restored_day = await session.get(Shift, day_id)
            restored_evening = await session.get(Shift, evening_id)
            assert restored_day is not None and restored_evening is not None
            assert as_local(restored_day.start_at, TZ) == as_local(original_day[0], TZ)
            assert as_local(restored_day.end_at, TZ) == as_local(original_day[1], TZ)
            assert as_local(restored_evening.start_at, TZ) == as_local(original_evening[0], TZ)
            assert as_local(restored_evening.end_at, TZ) == as_local(original_evening[1], TZ)
            assert await get_shift_times(session) == ShiftTimes()
        engine.dispose()

    asyncio.run(run())


def test_word_export_matches_official_two_group_layout_and_round_trips() -> None:
    shifts = []
    first = date(2026, 7, 12)
    for offset in range(37):
        duty_date = date.fromordinal(first.toordinal() + offset)
        shifts.append(
            SimpleNamespace(
                start_at=_utc(duty_date, time(13, 30)),
                end_at=_utc(duty_date, time(17, 0)),
                active=True,
                pharmacy=SimpleNamespace(name=f"نهارية {offset + 1}"),
            )
        )
        shifts.append(
            SimpleNamespace(
                start_at=_utc(duty_date, time(20, 30)),
                end_at=_utc(duty_date, time(23, 30)),
                active=True,
                pharmacy=SimpleNamespace(name=f"مسائية {offset + 1}"),
            )
        )

    data = build_official_word_schedule(shifts, TZ, ShiftTimes())
    document = Document(BytesIO(data))
    assert len(document.tables) == 1
    table = document.tables[0]
    assert len(table.columns) == 6
    assert len(table.rows) == 20
    assert any("جدول المنوبات لمدينة عامودة" in paragraph.text for paragraph in document.paragraphs)
    assert table.rows[1].cells[1].text == "نهارية 1"
    assert table.rows[1].cells[2].text == "مسائية 1"
    assert table.rows[1].cells[4].text == "نهارية 20"
    assert table.rows[1].cells[5].text == "مسائية 20"
    assert "الدوام المسائي" in document.paragraphs[-1].text

    parsed, warnings = parse_amuda_word_schedule(data)
    assert not warnings
    assert len(parsed) == 74
    assert parsed[0].pharmacy_name == "نهارية 1"
    assert parsed[-1].pharmacy_name == "مسائية 37"
