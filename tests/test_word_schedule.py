from __future__ import annotations

from datetime import date, time
from io import BytesIO

import pytest
from docx import Document

from app.services.word_schedule import parse_amuda_word_schedule


def _official_schedule_docx() -> bytes:
    document = Document()
    document.add_paragraph("جدول المنوبات لمدينة عامودة خلال شهر//7// تموز لعام 2026")
    table = document.add_table(rows=3, cols=6)
    headers = [
        "اليوم والتاريخ",
        "النهارية:\n١.٣٠-٥",
        "المساء:\n٨.٣٠-١١.٣٠",
        "اليوم والتاريخ",
        "النهارية:\n١.٣٠-٥",
        "المساء:\n٨.٣٠-١١.٣٠",
    ]
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value

    rows = [
        ["الاحد: 12/7/2026", "نور", "يوسف", "الجمعة: 31/7/2026", "سيدو", "حسن"],
        ["الاثنين: 13/7/2026", "بيمان", "رمضان", "السبت: 1/8/2026", "شيرين", "كاميران"],
    ]
    for row_index, values in enumerate(rows, start=1):
        for column_index, value in enumerate(values):
            table.rows[row_index].cells[column_index].text = value

    document.add_paragraph("الدوام المسائي:5:00 – 8:30")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_parses_two_side_by_side_periods_from_official_word_layout() -> None:
    parsed, warnings = parse_amuda_word_schedule(_official_schedule_docx())

    assert warnings == []
    assert len(parsed) == 8
    assert [(item.pharmacy_name, item.duty_date) for item in parsed[:4]] == [
        ("نور", date(2026, 7, 12)),
        ("يوسف", date(2026, 7, 12)),
        ("بيمان", date(2026, 7, 13)),
        ("رمضان", date(2026, 7, 13)),
    ]
    assert parsed[0].start_time == time(13, 30)
    assert parsed[0].end_time == time(17, 0)
    assert parsed[1].start_time == time(20, 30)
    assert parsed[1].end_time == time(23, 30)
    assert parsed[0].raw_data["period"] == "نهارية"
    assert parsed[1].raw_data["period"] == "مسائية"
    assert [item.row_number for item in parsed] == list(range(1, 9))


def test_keeps_dates_that_continue_into_the_next_month() -> None:
    parsed, _ = parse_amuda_word_schedule(_official_schedule_docx())
    august = [item for item in parsed if item.duty_date.month == 8]

    assert len(august) == 2
    assert {item.pharmacy_name for item in august} == {"شيرين", "كاميران"}
    assert all(item.duty_date == date(2026, 8, 1) for item in august)


def test_footer_regular_hours_are_not_imported_as_a_shift() -> None:
    parsed, _ = parse_amuda_word_schedule(_official_schedule_docx())

    assert all(item.pharmacy_name != "الدوام المسائي" for item in parsed)
    assert all(item.start_time != time(17, 0) or item.end_time != time(20, 30) for item in parsed)


def test_rejects_non_docx_input() -> None:
    with pytest.raises(ValueError, match="Word صالح"):
        parse_amuda_word_schedule(b"not-a-docx")
