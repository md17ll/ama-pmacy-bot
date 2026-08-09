from __future__ import annotations

import asyncio
import hashlib
from datetime import date, datetime, time, timezone
from io import BytesIO
from types import SimpleNamespace
from zoneinfo import ZoneInfo

from docx import Document

import app.handlers  # noqa: F401 - activates smart schedule patches
from app.services import smart_schedule_edit_patch as edit_patch
from app.services.shift_schedule_tools import ShiftTimes
from app.services.word_export import (
    TEMPLATE_SHA1,
    TEMPLATE_SIZE,
    build_official_word_schedule,
    template_bytes,
)


TZ = ZoneInfo("Asia/Damascus")


def _utc(day: date, value: time) -> datetime:
    return datetime.combine(day, value, tzinfo=TZ).astimezone(timezone.utc)


def _shift(day: date, value: time, name: str):
    return SimpleNamespace(
        start_at=_utc(day, value),
        end_at=_utc(day, time(17, 0) if value.hour < 18 else time(23, 30)),
        active=True,
        pharmacy=SimpleNamespace(name=name),
    )


def test_uploaded_docx_is_the_authoritative_word_template() -> None:
    raw = template_bytes()
    assert len(raw) == TEMPLATE_SIZE == 19840
    assert hashlib.sha1(raw).hexdigest() == TEMPLATE_SHA1 == "f42cc3b7780d27cc946c8411e6e66d15a429be02"
    source = Document(BytesIO(raw))
    assert len(source.tables) == 1
    assert len(source.tables[0].rows) == 20
    assert len(source.tables[0].columns) == 6
    assert source.tables[0].rows[0].cells[1].text == "النهارية: \n١.٣٠-٥"
    assert source.tables[0].rows[0].cells[2].text == "المساء:\n٨.٣٠-١١.٣٠"
    assert any(paragraph.text == "الدوام المسائي:5:00 – 8:30" for paragraph in source.paragraphs)


def test_word_export_puts_smart_day_and_evening_names_in_same_template_cells() -> None:
    shifts = [
        _shift(date(2026, 8, 18), time(13, 30), "صيدلية النهاري"),
        _shift(date(2026, 8, 18), time(20, 30), "صيدلية المسائي"),
        _shift(date(2026, 8, 19), time(13, 30), "صيدلية ثانية نهاري"),
        _shift(date(2026, 8, 19), time(20, 30), "صيدلية ثانية مسائي"),
    ]
    data = build_official_word_schedule(shifts, TZ, ShiftTimes())
    document = Document(BytesIO(data))
    table = document.tables[0]

    assert table.rows[0].cells[1].text == "النهارية: \n١.٣٠-٥"
    assert table.rows[0].cells[2].text == "المساء:\n٨.٣٠-١١.٣٠"
    assert table.rows[1].cells[0].text == "الثلاثاء: 18/8/2026"
    assert table.rows[1].cells[1].text == "صيدلية النهاري"
    assert table.rows[1].cells[2].text == "صيدلية المسائي"
    assert table.rows[2].cells[1].text == "صيدلية ثانية نهاري"
    assert table.rows[2].cells[2].text == "صيدلية ثانية مسائي"
    assert any(paragraph.text == "الدوام المسائي:5:00 – 8:30" for paragraph in document.paragraphs)
    assert any("شهر//8// آب لعام 2026" in paragraph.text for paragraph in document.paragraphs)


def test_word_export_repeats_same_template_for_more_than_38_dates() -> None:
    shifts = []
    first = date(2026, 8, 18)
    for offset in range(45):
        duty = date.fromordinal(first.toordinal() + offset)
        shifts.extend(
            [
                _shift(duty, time(13, 30), f"نهاري {offset + 1}"),
                _shift(duty, time(20, 30), f"مسائي {offset + 1}"),
            ]
        )
    data = build_official_word_schedule(shifts, TZ, ShiftTimes())
    document = Document(BytesIO(data))
    assert len(document.tables) == 2
    assert document.tables[0].rows[19].cells[4].text == "نهاري 38"
    assert document.tables[0].rows[19].cells[5].text == "مسائي 38"
    assert document.tables[1].rows[1].cells[1].text == "نهاري 39"
    assert document.tables[1].rows[1].cells[2].text == "مسائي 39"


def test_generated_choice_metadata_is_added_and_survives_manual_reroll(monkeypatch) -> None:
    duty = date(2026, 8, 18)
    start = _utc(duty, time(13, 30))

    async def fake_generate(*args, **kwargs):
        fixed = kwargs.get("fixed")
        pharmacy_id = fixed[(duty, "day")] if fixed else 7
        return (
            [
                {
                    "row_number": 1,
                    "raw_pharmacy_name": f"صيدلية {pharmacy_id}",
                    "matched_pharmacy_id": pharmacy_id,
                    "start_at": start,
                    "end_at": _utc(duty, time(17, 0)),
                    "confidence": 100.0,
                    "status": "ready",
                    "errors": [],
                    "raw_data": {"smart": True, "period": "day", "locked": bool(fixed)},
                }
            ],
            SimpleNamespace(),
        )

    monkeypatch.setattr(edit_patch, "_ORIGINAL_GENERATE_IMPORT_ROWS", fake_generate)

    rows, _ = asyncio.run(
        edit_patch.generate_import_rows(
            object(),
            start_date=duty,
            end_date=duty,
            timezone=TZ,
            times=ShiftTimes(),
        )
    )
    assert rows[0]["raw_data"]["generated_pharmacy_id"] == 7
    assert rows[0]["raw_data"]["generated_pharmacy_name"] == "صيدلية 7"

    fixed = edit_patch.FixedAssignments()
    fixed[(duty, "day")] = 9
    fixed.manual_keys.add((duty, "day"))
    fixed.generated_origins[(duty, "day")] = (7, "صيدلية 7")
    rows, _ = asyncio.run(
        edit_patch.generate_import_rows(
            object(),
            start_date=duty,
            end_date=duty,
            timezone=TZ,
            times=ShiftTimes(),
            fixed=fixed,
        )
    )
    data = rows[0]["raw_data"]
    assert rows[0]["matched_pharmacy_id"] == 9
    assert data["manual_override"] is True
    assert data["generated_pharmacy_id"] == 7
    assert data["generated_pharmacy_name"] == "صيدلية 7"
